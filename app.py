import streamlit as st
import google.generativeai as genai
import pandas as pd
from PyPDF2 import PdfReader
from docx import Document
import io
import time
import datetime
import plotly.express as px

# --- 1. הגדרות דף + CSS מלוטש ---
st.set_page_config(page_title="Summarizer Elite Pro 2026", page_icon="🧬", layout="wide")

def inject_custom_css():
    st.markdown("""
        <style>
        @keyframes gradientBg { 0% { background-position: 0% 50%; } 50% { background-position: 100% 50%; } 100% { background-position: 0% 50%; } }
        .stApp {
            background: linear-gradient(-45deg, #f0f4f8, #e2e8f0, #ffffff, #dbeafe);
            background-size: 400% 400%; animation: gradientBg 15s ease infinite;
        }
        .stButton>button {
            width: 100%; border-radius: 12px; height: 3.5em;
            background: linear-gradient(135deg, #2563eb 0%, #1e3a8a 100%);
            color: white; font-weight: bold; border: none;
            transition: 0.3s; box-shadow: 0 4px 15px rgba(37,99,235,0.2);
        }
        .stButton>button:hover { transform: scale(1.03); box-shadow: 0 8px 25px rgba(37,99,235,0.3); }
        .result-card {
            background: rgba(255,255,255,0.85); backdrop-filter: blur(12px);
            padding: 25px; border-radius: 16px; box-shadow: 0 10px 35px rgba(0,0,0,0.06);
            margin-bottom: 20px; border: 1px solid rgba(255,255,255,0.4);
        }
        .stats-box { background: white; padding: 18px; border-radius: 12px; text-align: center;
            border: 1px solid #e0e0e0; box-shadow: 0 4px 15px rgba(0,0,0,0.03); }
        .rtl-container { direction: rtl; text-align: right; }
        .ltr-container { direction: ltr; text-align: left; }
        .copy-btn { background: #2563eb; color: white; border: none; padding: 8px 16px;
            border-radius: 8px; cursor: pointer; font-weight: bold; }
        </style>
    """, unsafe_allow_html=True)
inject_custom_css()

# --- 2. Session State ---
if 'history' not in st.session_state:
    st.session_state.history = []
if 'analysis_results' not in st.session_state:
    st.session_state.analysis_results = []

# --- 3. פונקציות עזר משופרות ---
def get_text_from_file(uploaded_file):
    if not uploaded_file:
        return ""
    if uploaded_file.type == "application/pdf":
        reader = PdfReader(uploaded_file)
        return "".join([page.extract_text() or "" for page in reader.pages])
    elif uploaded_file.type == "text/plain":
        return uploaded_file.read().decode("utf-8")
    elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        doc = Document(uploaded_file)
        return "\n".join([para.text for para in doc.paragraphs])
    return ""

def create_docx(text, model_name):
    doc = Document()
    doc.add_heading(f'Summary Report - {model_name}', 0)
    doc.add_paragraph(f"Date: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M')}")
    doc.add_heading('Summary Content:', level=1)
    doc.add_paragraph(text)
    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()

def perform_sentiment_analysis(text):
    positive = ['good', 'great', 'excellent', 'success', 'טוב', 'מעולה', 'הצלחה', 'נפלא']
    negative = ['bad', 'poor', 'error', 'failure', 'רע', 'גרוע', 'כישלון', 'בעיה']
    pos = sum(1 for w in positive if w in text.lower())
    neg = sum(1 for w in negative if w in text.lower())
    if pos > neg: return "Positive 😊", "#28a745"
    if neg > pos: return "Negative 😟", "#dc3545"
    return "Neutral 😐", "#6c757d"

# --- 4. ממשק ---
st.title("Summarizer Elite Pro v6.0 🧬 – מנוע AI משולב")

with st.sidebar:
    st.header("⚙️ פאנל ניהול")
    
    model_list = [
        'models/gemini-2.5-pro',
        'models/gemini-2.5-flash',
        'models/gemini-2.0-flash',
        'models/gemini-2.5-flash-lite',
        'models/gemini-1.5-pro'
    ]
    
    selected_models = st.multiselect(
        "בחר מודלים להרצה (לפי סדר עדיפות):",
        model_list,
        default=[model_list[0], model_list[1]]
    )
    
    lang = st.selectbox("שפת הסיכום:", ["Hebrew", "English", "Spanish", "Russian"])
    detail = st.select_slider("רמת פירוט:", ["תמציתי", "מאוזן", "מפורט"])
    detail_map = {"תמציתי": "concise", "מאוזן": "balanced", "מפורט": "detailed"}
    
    st.markdown("---")
    api_key = st.text_input("מפתח Gemini API:", type="password", placeholder="AIza...")
    final_api_key = api_key or st.secrets.get("GEMINI_API_KEY", "")

# טאבים
tab1, tab2, tab3 = st.tabs(["🚀 מרכז עיבוד", "📜 היסטוריה", "📊 סטטיסטיקה"])

with tab1:
    col_in, col_stats = st.columns([2, 1])
    
    with col_in:
        uploaded_file = st.file_uploader(
            "העלה מסמך (PDF / TXT / DOCX):",
            type=["pdf", "txt", "docx"]
        )
        raw_text = st.text_area("או הדבק טקסט כאן:", height=220)
        
        content = get_text_from_file(uploaded_file) if uploaded_file else raw_text
    
    with col_stats:
        if content:
            words = len(content.split())
            chars = len(content)
            st.markdown(f'<div class="stats-box"><h3>מילים</h3><h2>{words:,}</h2></div>', unsafe_allow_html=True)
            st.markdown(f'<div class="stats-box"><h3>תווים</h3><h2>{chars:,}</h2></div>', unsafe_allow_html=True)
            sentiment, color = perform_sentiment_analysis(content)
            st.markdown(f'<div class="stats-box" style="border-top: 6px solid {color}"><h3>סנטימנט</h3><h2 style="color:{color}">{sentiment}</h2></div>', unsafe_allow_html=True)
    
    if st.button("🚀 הפעל מנוע AI משולב ✨", type="primary"):
        if not content.strip():
            st.error("⚠️ יש להזין תוכן")
        elif not final_api_key:
            st.error("⚠️ חובה להזין מפתח Gemini API")
        else:
            genai.configure(api_key=final_api_key)
            st.session_state.analysis_results = []
            
            detail_eng = detail_map[detail]
            
            for m_name in selected_models:
                with st.spinner(f"מעבד עם {m_name}..."):
                    try:
                        model = genai.GenerativeModel(
                            m_name,
                            safety_settings=[
                                {"category": "HARM_CATEGORY_HARASSMENT", "threshold": "BLOCK_MEDIUM_AND_ABOVE"},
                                {"category": "HARM_CATEGORY_HATE_SPEECH", "threshold": "BLOCK_MEDIUM_AND_ABOVE"}
                            ]
                        )
                        prompt = f"""
You are an expert document summarizer.
Summarize the following text in {lang} language at {detail_eng} level of detail.
Structure the summary with:
• Main points
• Key insights
• Clear conclusion

Text:
{content}
"""
                        start = time.time()
                        response = model.generate_content(prompt)
                        end = time.time()
                        
                        res_data = {
                            "model": m_name,
                            "text": response.text,
                            "time": round(end - start, 2),
                            "date": datetime.datetime.now().strftime("%H:%M")
                        }
                        st.session_state.analysis_results.append(res_data)
                        st.session_state.history.append(res_data)
                        st.success(f"✅ {m_name} הצליח")
                        
                    except Exception as e:
                        st.warning(f"⚠️ {m_name} נכשל: {str(e)[:120]}")
                        continue
            
            if not st.session_state.analysis_results:
                st.error("כל המודלים נכשלו. בדוק מפתח API או נסה שוב.")

    # הצגת תוצאות
    if st.session_state.analysis_results:
        alignment = "rtl-container" if lang == "Hebrew" else "ltr-container"
        for res in st.session_state.analysis_results:
            st.markdown(f'''
                <div class="result-card">
                    <div class="{alignment}">
                        <h3 style="color:#1e3a8a;">🤖 {res['model']}</h3>
                        <p style="font-size:1.05rem; line-height:1.7;">{res['text']}</p>
                        <p style="color:gray; font-size:0.85rem;">⏱️ {res['time']} שניות</p>
                    </div>
                </div>
            ''', unsafe_allow_html=True)
            
            c1, c2, c3 = st.columns([1,1,1])
            c1.download_button("📥 TXT", res['text'], f"{res['model'].replace('/','_')}.txt", key=f"t_{res['model']}")
            c2.download_button("📄 Word", create_docx(res['text'], res['model']), f"{res['model'].replace('/','_')}.docx", key=f"d_{res['model']}")
            c3.button("📋 העתק", key=f"copy_{res['model']}", 
                      on_click=lambda t=res['text']: st.toast("✅ הועתק ללוח!", icon="📋") or st.session_state.update({"clipboard": t}))

with tab2:
    if st.session_state.history:
        st.dataframe(pd.DataFrame(st.session_state.history), use_container_width=True)
        if st.button("🗑️ נקה היסטוריה"):
            st.session_state.history = []
            st.rerun()

with tab3:
    if st.session_state.history:
        df = pd.DataFrame(st.session_state.history)
        fig = px.bar(df, x="model", y="time", title="זמני עיבוד לפי מודל", color="model", text_auto=True)
        st.plotly_chart(fig, use_container_width=True)

st.caption("✅ Summarizer Elite Pro 2026 – קוד מלוטש ומשופר במיוחד עבורך")
